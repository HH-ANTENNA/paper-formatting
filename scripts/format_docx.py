#!/usr/bin/env python3
"""
format_docx.py — Format a Word document according to the default paper template.

Phases: A=Page+Body, B=FrontMatter, C=Headings, D=Tables+Figures, E=References
Run with --phase all (default) or specify individual phases.
"""
import json
import sys
import os
import re
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement


# ═══════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════

def is_chinese_char(ch):
    cp = ord(ch)
    return (0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF or
            0x20000 <= cp <= 0x2A6DF or 0xF900 <= cp <= 0xFAFF or
            0x3000 <= cp <= 0x303F or 0xFF00 <= cp <= 0xFFEF)


def text_is_chinese(text):
    if not text:
        return False
    chinese_count = sum(1 for c in text if is_chinese_char(c))
    return chinese_count > len(text) * 0.3


def set_run_font(run, chinese_font, english_font, size_pt, bold=None, color_rgb=None):
    run.font.size = Pt(size_pt)
    run.font.name = english_font
    if bold is not None:
        run.bold = bold
    if color_rgb is not None:
        run.font.color.rgb = color_rgb
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), chinese_font)
    rFonts.set(qn('w:ascii'), english_font)
    rFonts.set(qn('w:hAnsi'), english_font)
    rFonts.set(qn('w:cs'), english_font)


def set_paragraph_spacing(paragraph, line_spacing=1.5, space_before=0, space_after=0,
                          first_line_indent_cm=0.74, alignment=None):
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)
    if alignment is not None:
        paragraph.alignment = alignment


# English → Chinese punctuation mapping
_EN_TO_CN_PUNCT = {
    ',': '，', ':': '：', ';': '；', '?': '？', '!': '！',
    '(': '（', ')': '）',
}

_PRESERVE_PATTERNS = [
    re.compile(r'https?://[a-zA-Z0-9._~:/?#\[\]@!$&()*+;=%\-]+'),
    re.compile(r'ftp://[a-zA-Z0-9._~:/?#\[\]@!$&()*+;=%\-]+'),
    re.compile(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}'),
    re.compile(r'\(\d+(?:\.\d+)?\)'),
    re.compile(r'(?<!\d)v?\d+\.\d+(?:\.\d+)*(?:\-[a-zA-Z0-9]+)?(?![.\d])'),
    re.compile(r'(?<!\d)\d+\.\d+(?![.\d])'),
    re.compile(r'(?:^|(?<=\s))\d+(?:\.\d+)*\.(?=\s|$)'),
    re.compile(
        r'(?:^|(?<=[\s(（])|(?<=[一-鿿]))'
        r'(?:i\.e\.|e\.g\.|etc\.|vs\.|et\s+al\.|ca\.|approx\.)'
        r'(?:[,;]?\s*(?:i\.e\.|e\.g\.|etc\.|vs\.|et\s+al\.|ca\.|approx\.))*'
        r'(?:[,;]?)',
        re.IGNORECASE
    ),
]


def _protect_patterns(text):
    protected = {}
    counter = [0]
    def replace(match):
        placeholder = f'\x00PROTECT\x00{counter[0]}\x00'
        protected[placeholder] = match.group(0)
        counter[0] += 1
        return placeholder
    for pattern in _PRESERVE_PATTERNS:
        text = pattern.sub(replace, text)
    return text, protected


def _restore_patterns(text, protected):
    for placeholder, original in protected.items():
        text = text.replace(placeholder, original)
    return text


def _convert_periods_by_context(text, full_stop):
    chars = list(text)
    n = len(chars)

    def _find_real_prev(j):
        while j >= 0:
            ch = chars[j]
            if ch == '\x00':
                j -= 1
                while j >= 0 and chars[j] != '\x00':
                    j -= 1
                if j >= 0:
                    j -= 1
                continue
            cp = ord(ch)
            if cp in (0x29, 0x5D, 0x7D, 0x3E):
                j -= 1; continue
            if 0x3000 <= cp <= 0x303F or 0xFF00 <= cp <= 0xFFEF:
                j -= 1; continue
            return ch
        return ''

    for i, ch in enumerate(chars):
        if ch != '.':
            continue
        next_char = chars[i + 1] if i + 1 < n else ''
        prev_char = chars[i - 1] if i > 0 else ''
        if prev_char.isdigit() and next_char.isdigit():
            continue
        if prev_char.isalpha() and prev_char.isascii():
            continue
        if is_chinese_char(prev_char):
            chars[i] = full_stop; continue
        if prev_char.isdigit():
            chars[i] = full_stop; continue
        cp = ord(prev_char) if prev_char else 0
        if (cp in (0x29, 0x5D, 0x7D, 0x3E)
                or 0x3000 <= cp <= 0x303F
                or 0xFF00 <= cp <= 0xFFEF
                or prev_char == '\x00'):
            real_prev = _find_real_prev(i - 1)
            if real_prev and (is_chinese_char(real_prev) or real_prev.isdigit()):
                chars[i] = full_stop
    return ''.join(chars)


def _convert_other_punct_by_context(text):
    chars = list(text)
    for i, ch in enumerate(chars):
        if ch not in _EN_TO_CN_PUNCT:
            continue
        prev_char = chars[i - 1] if i > 0 else ''
        next_char = chars[i + 1] if i + 1 < len(chars) else ''
        if is_chinese_char(prev_char) or is_chinese_char(next_char):
            chars[i] = _EN_TO_CN_PUNCT[ch]
    return ''.join(chars)


def fix_chinese_punctuation(text, full_stop=None):
    if not text:
        return text
    if full_stop is None:
        full_stop = '。'
    if full_stop == '．':
        text = text.replace('。', '．')
    text, protected = _protect_patterns(text)
    text = _convert_periods_by_context(text, full_stop)
    text = _convert_other_punct_by_context(text)
    text = _restore_patterns(text, protected)
    text = _convert_periods_by_context(text, full_stop)
    return text


# ═══════════════════════════════════════════════════════════
# Paragraph type detection
# ═══════════════════════════════════════════════════════════

def detect_paragraph_type(paragraph, rules=None):
    text = paragraph.text.strip()
    if not text:
        return "empty"

    style_name = paragraph.style.name if paragraph.style else ''

    # Style-based detection
    if style_name == 'Heading 1' or style_name == 'heading 1':
        return "heading_l1"
    if style_name == 'Heading 2' or style_name == 'heading 2':
        return "heading_l2"
    if style_name == 'Heading 3' or style_name == 'heading 3':
        return "heading_l3"

    bilingual = rules.get("bilingual_required", False) if rules else False

    # Abstract / Keywords (支持两种格式: 【摘要】 和 摘要：)
    if re.match(r'^【?\s*摘\s*要\s*[】》：:]', text):
        return "abstract_chinese_heading"
    if re.match(r'^Abstract[：:]?', text, re.IGNORECASE):
        return "abstract_english_heading"
    if re.match(r'^【?\s*关键[词字]\s*[】》：:]', text):
        return "keywords_chinese"
    if re.match(r'^Key words?[：:]', text, re.IGNORECASE):
        return "keywords_english"

    # Numbered headings: "1  XXXX", "1.1  XXXX", "1.1.1  XXXX"
    heading_match = re.match(r'^(\d+(?:\.\d+)*)\s{1,2}\S', text)
    if heading_match:
        depth = heading_match.group(1).count('.') + 1
        return f"heading_l{depth}" if depth <= 3 else "heading_l3"

    # Unnumbered headings (引言, 绪论, etc.)
    if text in ['引言', '绪论', '结语', '结论', '致谢']:
        return "unnumbered_heading"

    # Reference section
    if text in ['参考文献', 'References'] or text.startswith('参考文献'):
        return "reference_heading"

    # Figure / Table captions
    if re.match(r'^图\s*\d+', text):
        return "figure_caption_chinese"
    if re.match(r'^Fig\.?\s*\d+', text, re.IGNORECASE):
        return "figure_caption_english"
    if re.match(r'^表\s*\d+', text):
        return "table_caption_chinese"
    if re.match(r'^Table\s+\d+', text, re.IGNORECASE):
        return "table_caption_english"

    # Reference entries
    if re.match(r'^\[\d+\]', text):
        return "reference_entry"

    # Author / affiliation (journal mode)
    if re.match(r'^作者\d+', text):
        return "author_chinese"
    if re.match(r'^Author\s+\d+', text, re.IGNORECASE):
        return "author_english"
    if re.match(r'^\d+\)', text):
        return "affiliation"

    return "body"


# ═══════════════════════════════════════════════════════════
# Phase A: Page + Body
# ═══════════════════════════════════════════════════════════

def phase_a_page_and_body(doc, rules):
    changes = []

    page_rules = rules.get("extracted", {}).get("page", rules.get("page", {}))
    for section in doc.sections:
        if page_rules.get("width_cm"):
            section.page_width = Cm(page_rules["width_cm"])
            section.page_height = Cm(page_rules["height_cm"])
            changes.append(f"Page size: {page_rules['width_cm']}x{page_rules['height_cm']}cm")
        if page_rules.get("top_margin_cm"):
            section.top_margin = Cm(page_rules["top_margin_cm"])
            section.bottom_margin = Cm(page_rules["bottom_margin_cm"])
            section.left_margin = Cm(page_rules["left_margin_cm"])
            section.right_margin = Cm(page_rules["right_margin_cm"])
            changes.append(f"Margins: T={page_rules['top_margin_cm']} B={page_rules['bottom_margin_cm']} "
                          f"L={page_rules['left_margin_cm']} R={page_rules['right_margin_cm']}cm")

    body = rules.get("body_font", {})
    chinese_font = body.get("chinese", "宋体")
    english_font = body.get("english", "Times New Roman")
    size_pt = body.get("size_pt", 11)
    line_spacing = rules.get("line_spacing", 1.15)
    indent_cm = rules.get("first_line_indent_cm", 0.74)

    punctuation_rules = rules.get("punctuation", {})
    full_stop = punctuation_rules.get("chinese_full_stop", "。")

    for p in doc.paragraphs:
        ptype = detect_paragraph_type(p, rules)
        if ptype == "empty":
            continue
        text = p.text.strip()
        if not text:
            continue

        # Fix punctuation in Chinese body / abstract / keywords
        if ptype in ("body", "abstract_chinese_heading", "keywords_chinese"):
            if text_is_chinese(text) and p.runs:
                for run in p.runs:
                    run.text = fix_chinese_punctuation(run.text, full_stop)

        # Apply body formatting
        if ptype == "body":
            for run in p.runs:
                set_run_font(run, chinese_font, english_font, size_pt,
                            color_rgb=RGBColor(0, 0, 0))
            set_paragraph_spacing(p, line_spacing, first_line_indent_cm=indent_cm,
                                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    changes.append(f"Body: {chinese_font}/{english_font} {size_pt}pt, "
                   f"line {line_spacing}, indent {indent_cm}cm")
    return changes


# ═══════════════════════════════════════════════════════════
# Phase B: Front Matter
# ═══════════════════════════════════════════════════════════

def phase_b_front_matter(doc, rules):
    changes = []

    title_font = rules.get("title", {})
    abstract_font = rules.get("abstract", {})
    abstract_label_font = rules.get("abstract_label", abstract_font)
    kw_font = rules.get("keywords", {})
    kw_label_font = rules.get("keywords_label", kw_font)
    bilingual = rules.get("bilingual_required", False)

    title_found = False
    in_abstract = False

    for p in doc.paragraphs:
        ptype = detect_paragraph_type(p, rules)
        text = p.text.strip()
        if not text:
            in_abstract = False
            continue

        if ptype == "empty":
            continue

        # ── Title ──
        if not title_found and ptype == "body":
            pf = p.paragraph_format
            is_centered = (p.alignment == WD_ALIGN_PARAGRAPH.CENTER or
                          pf.alignment == WD_ALIGN_PARAGRAPH.CENTER)
            if is_centered or (len(text) < 80 and not re.match(r'^\d', text)):
                for run in p.runs:
                    set_run_font(run,
                                title_font.get("chinese", "宋体"),
                                title_font.get("english", "Times New Roman"),
                                title_font.get("size_pt", 18),
                                bold=title_font.get("bold", True),
                                color_rgb=RGBColor(0, 0, 0))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(p, rules.get("line_spacing", 1.15),
                                     space_before=12, space_after=12,
                                     first_line_indent_cm=0)
                title_found = True
                changes.append(f"Title: {text[:50]}")
                continue

        # ── Abstract (inline: "【摘要】content" or "摘要：content") ──
        if ptype == "abstract_chinese_heading":
            in_abstract = True
            if p.runs:
                # Normalize: 【摘要】 → keep, 摘要： → 【摘要】
                p.runs[0].text = re.sub(r'^摘\s*要[：:]', '【摘要】', p.runs[0].text)
                p.runs[0].text = re.sub(r'^【?\s*摘\s*要\s*】?\s*', '【摘要】', p.runs[0].text)
                for j, run in enumerate(p.runs):
                    if j == 0:
                        set_run_font(run,
                                    abstract_label_font.get("chinese", "楷体"),
                                    abstract_label_font.get("english", "Times New Roman"),
                                    abstract_label_font.get("size_pt", 12),
                                    bold=False)
                    else:
                        set_run_font(run,
                                    abstract_font.get("chinese", "楷体"),
                                    abstract_font.get("english", "Times New Roman"),
                                    abstract_font.get("size_pt", 10.5),
                                    bold=False)
            set_paragraph_spacing(p, rules.get("line_spacing", 1.15),
                                 first_line_indent_cm=0.74)
            changes.append(f"Abstract: {text[:50]}")
            continue

        # ── English Abstract (when bilingual enabled) ──
        if ptype == "abstract_english_heading" and bilingual:
            for run in p.runs:
                set_run_font(run, abstract_font.get("english", "Times New Roman"),
                            abstract_font.get("english", "Times New Roman"),
                            abstract_font.get("size_pt", 10.5), bold=True)
            changes.append(f"English abstract: {text[:50]}")
            continue

        # ── Abstract body (continuation paragraphs between abstract and keywords) ──
        if in_abstract and ptype == "body" and not title_found:
            for run in p.runs:
                set_run_font(run, abstract_font.get("chinese", "楷体"),
                            abstract_font.get("english", "Times New Roman"),
                            abstract_font.get("size_pt", 10.5))
            set_paragraph_spacing(p, rules.get("line_spacing", 1.15),
                                 first_line_indent_cm=0.74)
            changes.append(f"Abstract body: {text[:50]}")
            continue

        # ── Keywords (inline: "【关键词】item1；item2" or "关键词：item1；item2") ──
        if ptype == "keywords_chinese":
            in_abstract = False
            if p.runs:
                # Normalize: 【关键词】 → keep, 关键词： → 【关键词】
                p.runs[0].text = re.sub(r'^关键[词字][：:]', '【关键词】', p.runs[0].text)
                p.runs[0].text = re.sub(r'^【?\s*关键[词字]\s*】?\s*', '【关键词】', p.runs[0].text)
                for j, run in enumerate(p.runs):
                    if j == 0:
                        set_run_font(run,
                                    kw_label_font.get("chinese", "楷体"),
                                    kw_label_font.get("english", "Times New Roman"),
                                    kw_label_font.get("size_pt", 12),
                                    bold=kw_label_font.get("bold", True))
                    else:
                        set_run_font(run,
                                    kw_font.get("chinese", "楷体"),
                                    kw_font.get("english", "Times New Roman"),
                                    kw_font.get("size_pt", 10.5))
                # Fix separator
                for run in p.runs:
                    run.text = run.text.replace(',', '；').replace('，', '；')
            set_paragraph_spacing(p, rules.get("line_spacing", 1.15),
                                 first_line_indent_cm=0)
            changes.append(f"Keywords: {text[:50]}")
            continue

        # ── English Keywords (when bilingual enabled) ──
        if ptype == "keywords_english" and bilingual:
            for run in p.runs:
                set_run_font(run, kw_font.get("english", "Times New Roman"),
                            kw_font.get("english", "Times New Roman"),
                            kw_font.get("size_pt", 10.5))
            changes.append(f"English keywords: {text[:50]}")
            continue

        # ── Author (journal mode) ──
        if ptype in ("author_chinese", "author_english"):
            author_font = rules.get("author", {})
            for run in p.runs:
                set_run_font(run,
                            author_font.get("chinese", "楷体"),
                            author_font.get("english", "Times New Roman"),
                            author_font.get("size_pt", 14))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, rules.get("line_spacing", 1.15),
                                 first_line_indent_cm=0)
            changes.append(f"Author: {text[:50]}")
            continue

    return changes


# ═══════════════════════════════════════════════════════════
# Phase C: Headings
# ═══════════════════════════════════════════════════════════

def phase_c_headings(doc, rules):
    """Format heading hierarchy: L1, L2, L3.

    Applies Word built-in Heading 1/2/3 styles (which enables collapse/expand
    in the navigation pane), then overrides font formatting to match template.
    """
    changes = []

    l1 = rules.get("heading_l1", {"chinese": "宋体", "english": "Times New Roman", "size_pt": 14, "bold": True})
    l2 = rules.get("heading_l2", {"chinese": "宋体", "english": "Times New Roman", "size_pt": 15, "bold": True})
    l3 = rules.get("heading_l3", {"chinese": "楷体", "english": "Times New Roman", "size_pt": 14, "bold": False})

    heading_config = {
        "heading_l1": ("Heading 1", l1),
        "heading_l2": ("Heading 2", l2),
        "heading_l3": ("Heading 3", l3),
        "unnumbered_heading": ("Heading 1", l1),
    }

    line_spacing = rules.get("line_spacing", 1.15)

    for p in doc.paragraphs:
        ptype = detect_paragraph_type(p, rules)
        if ptype not in heading_config:
            continue

        style_name, font_config = heading_config[ptype]

        # Apply built-in Heading style (enables collapse/expand in Word)
        try:
            p.style = doc.styles[style_name]
        except KeyError:
            pass  # Style not found, skip but still apply font formatting below

        for run in p.runs:
            set_run_font(run,
                        font_config["chinese"], font_config["english"],
                        font_config["size_pt"],
                        bold=font_config.get("bold"),
                        color_rgb=RGBColor(0, 0, 0))  # Override theme color from Heading style

        set_paragraph_spacing(p, line_spacing, space_before=12, space_after=6,
                              first_line_indent_cm=0, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        changes.append(f"{ptype} → {style_name} ({font_config.get('chinese','?')} "
                       f"{font_config.get('size_pt','?')}pt"
                       f"{' Bold' if font_config.get('bold') else ''}): "
                       f"{p.text.strip()[:50]}")

    return changes


# ═══════════════════════════════════════════════════════════
# Phase D: Tables + Figures
# ═══════════════════════════════════════════════════════════

def phase_d_tables_figures(doc, rules):
    changes = []
    caption_font = rules.get("caption", {"chinese": "宋体", "english": "Times New Roman", "size_pt": 9})
    bilingual = rules.get("bilingual_required", False)

    for p in doc.paragraphs:
        ptype = detect_paragraph_type(p, rules)
        if ptype in ("table_caption_chinese", "table_caption_english",
                     "figure_caption_chinese", "figure_caption_english"):

            if not bilingual and "english" in ptype:
                continue

            for run in p.runs:
                set_run_font(run, caption_font["chinese"], caption_font["english"],
                            caption_font["size_pt"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, 1.0, first_line_indent_cm=0)
            changes.append(f"Caption: {p.text.strip()[:50]}")

    for table in doc.tables:
        apply_three_line_table(table)
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, caption_font["chinese"], caption_font["english"],
                                    caption_font["size_pt"])
        changes.append(f"Table → three-line ({len(table.rows)} rows)")

    return changes


def apply_three_line_table(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    for e in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(e)
    tblPr.append(tblBorders)

    num_rows = len(table.rows)
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')

            if ri == 0:
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), '12')
                top.set(qn('w:color'), '000000')
                tcBorders.append(top)
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:color'), '000000')
                tcBorders.append(bottom)
            elif ri == num_rows - 1:
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '12')
                bottom.set(qn('w:color'), '000000')
                tcBorders.append(bottom)

            for side in ['left', 'right']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)

            for eb in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(eb)
            tcPr.append(tcBorders)


# ═══════════════════════════════════════════════════════════
# Phase E: References
# ═══════════════════════════════════════════════════════════

def phase_e_references(doc, rules):
    changes = []

    ref_title_font = rules.get("reference_title",
                               {"chinese": "楷体", "english": "Times New Roman", "size_pt": 10.5})
    ref_body_font = rules.get("reference_body",
                              {"chinese": "宋体", "english": "Times New Roman", "size_pt": 9})
    hanging_cm = rules.get("reference_hanging_cm", 0.74)

    in_refs = False
    for p in doc.paragraphs:
        ptype = detect_paragraph_type(p, rules)
        text = p.text.strip()
        if not text:
            continue

        # Reference section header
        if ptype == "reference_heading" or text == '参考文献' or text.startswith('参考文献'):
            in_refs = True
            for run in p.runs:
                set_run_font(run, ref_title_font["chinese"], ref_title_font["english"],
                            ref_title_font["size_pt"],
                            bold=ref_title_font.get("bold", False))
            set_paragraph_spacing(p, rules.get("line_spacing", 1.15),
                                 space_before=12, first_line_indent_cm=0)
            changes.append("Reference title formatted")
            continue

        if in_refs and ptype == "reference_entry":
            for run in p.runs:
                set_run_font(run, ref_body_font["chinese"], ref_body_font["english"],
                            ref_body_font["size_pt"])
            pf = p.paragraph_format
            pf.first_line_indent = Cm(-hanging_cm)
            pf.left_indent = Cm(hanging_cm)
            changes.append(f"Reference: {text[:50]}")
            continue

        # Fix "1." → "[1]"
        if in_refs and re.match(r'^\d+\.', text):
            for run in p.runs:
                set_run_font(run, ref_body_font["chinese"], ref_body_font["english"],
                            ref_body_font["size_pt"])
            if p.runs:
                p.runs[0].text = re.sub(r'^(\d+)\.', r'[\1]', p.runs[0].text)
            pf = p.paragraph_format
            pf.first_line_indent = Cm(-hanging_cm)
            pf.left_indent = Cm(hanging_cm)
            changes.append(f"Reference fixed: {text[:50]}")

    return changes


# ═══════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════

PHASES = {
    "A": ("Page + Body", phase_a_page_and_body),
    "B": ("Front Matter", phase_b_front_matter),
    "C": ("Headings", phase_c_headings),
    "D": ("Tables + Figures", phase_d_tables_figures),
    "E": ("References", phase_e_references),
}


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Format academic paper (.docx)")
    parser.add_argument("input", help="Input .docx file")
    parser.add_argument("rules", help="Rules JSON file (from extract_rules.py)")
    parser.add_argument("--output", "-o", required=True, help="Output .docx file")
    parser.add_argument("--phase", default="all",
                       help="Phases: A,B,C,D,E or 'all' (default)")
    args = parser.parse_args()

    with open(args.rules, 'r', encoding='utf-8') as f:
        rules = json.load(f)

    doc = Document(args.input)

    print(f"Template: {rules.get('template_name', 'Custom')}")
    print(f"Body: {rules.get('body_font', {})}")
    print(f"Bilingual: {rules.get('bilingual_required', False)}")
    print()

    if args.phase == "all":
        phases_to_run = ["A", "B", "C", "D", "E"]
    else:
        phases_to_run = [p.strip() for p in args.phase.split(",")]

    all_changes = []
    for phase_id in phases_to_run:
        if phase_id not in PHASES:
            print(f"Unknown phase: {phase_id}")
            continue
        name, func = PHASES[phase_id]
        print(f"{'='*60}")
        print(f"Phase {phase_id}: {name}")
        print(f"{'='*60}")
        changes = func(doc, rules)
        for c in changes:
            print(f"  [OK] {c}")
        all_changes.extend(changes)

    doc.save(args.output)
    print(f"\n{'='*60}")
    print(f"Saved: {args.output}")
    print(f"Total changes: {len(all_changes)}")


if __name__ == '__main__':
    main()
