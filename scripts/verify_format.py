#!/usr/bin/env python3
"""
verify_format.py — Check a formatted document against the default template rules.

Reports: ✅ Passed  ⚠️ Warnings  ❌ Manual fixes needed
"""
import json
import sys
import re
from docx import Document
from docx.shared import Cm, Pt


def is_chinese_char(ch):
    if not ch:
        return False
    cp = ord(ch)
    return (0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF or
            0x20000 <= cp <= 0x2A6DF or 0xF900 <= cp <= 0xFAFF or
            0x3000 <= cp <= 0x303F or 0xFF00 <= cp <= 0xFFEF)


_EN_PUNCT_IN_CN_WARNING = {
    ',': ('，(U+FF0C)', 'English comma "," in Chinese text'),
    ';': ('；(U+FF1B)', 'English semicolon ";" in Chinese text'),
    ':': ('：(U+FF1A)', 'English colon ":" in Chinese text after Chinese char'),
    '?': ('？(U+FF1F)', 'English question mark "?" in Chinese text'),
    '!': ('！(U+FF01)', 'English exclamation mark "!" in Chinese text'),
    '(': ('（(U+FF08)', 'English left parenthesis "(" in Chinese text'),
    ')': ('）(U+FF09)', 'English right parenthesis ")" in Chinese text'),
}


def check_document(doc_path, rules):
    doc = Document(doc_path)
    report = {"passed": [], "warnings": [], "manual_fixes": []}

    bilingual = rules.get("bilingual_required", False)
    template_name = rules.get("template_name", "Unknown")

    page_rules = rules.get("extracted", {}).get("page", rules.get("page", {}))
    expected_margins = {
        "top": page_rules.get("top_margin_cm"),
        "bottom": page_rules.get("bottom_margin_cm"),
        "left": page_rules.get("left_margin_cm"),
        "right": page_rules.get("right_margin_cm"),
    }

    # ── Page size & margins ──
    for i, section in enumerate(doc.sections):
        w_cm = section.page_width / 360000
        h_cm = section.page_height / 360000
        if abs(w_cm - 21.0) > 0.5 or abs(h_cm - 29.7) > 0.5:
            report["warnings"].append(f"S{i}: page {w_cm:.1f}x{h_cm:.1f}cm (expected A4)")
        else:
            report["passed"].append(f"S{i}: A4 ({w_cm:.1f}x{h_cm:.1f}cm)")

        if expected_margins["top"]:
            t_cm = section.top_margin / 360000
            b_cm = section.bottom_margin / 360000
            l_cm = section.left_margin / 360000
            r_cm = section.right_margin / 360000
            margin_ok = True
            for name, exp, act in [
                ("top", expected_margins["top"], t_cm),
                ("bottom", expected_margins["bottom"], b_cm),
                ("left", expected_margins["left"], l_cm),
                ("right", expected_margins["right"], r_cm),
            ]:
                if abs(exp - act) > 1.0:
                    report["warnings"].append(f"S{i}: {name} margin {act:.1f}cm (expected {exp}cm)")
                    margin_ok = False
            if margin_ok:
                report["passed"].append(f"S{i}: margins OK")

    # ── Paragraph checks ──
    has_english_abstract = False
    has_english_keywords = False
    has_chinese_refs = False
    has_english_refs = False
    punctuation_issues = []
    full_stop_issues = []
    has_abstract_inline = False
    has_keywords_inline = False

    punctuation_rules = rules.get("punctuation", {})
    expected_full_stop = punctuation_rules.get("chinese_full_stop", "。")
    full_stop_name = "．(U+FF0E)" if expected_full_stop == "．" else "。(U+3002)"

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        # Bilingual checks
        if text.lower().startswith('abstract'):
            has_english_abstract = True
        if text.lower().startswith('key words') or text.lower().startswith('keywords'):
            has_english_keywords = True

        # Reference analysis
        if re.match(r'^\[\d+\]', text):
            if any(is_chinese_char(c) for c in text):
                has_chinese_refs = True
            else:
                has_english_refs = True

        # Abstract / Keywords format (支持 【摘要】 和 摘要：两种)
        if re.search(r'摘要[】》：:]', text):
            has_abstract_inline = True
        if re.search(r'关键[词字][】》：:]', text):
            has_keywords_inline = True
            if '；' in text:
                pass  # good
            elif ',' in text or '，' in text:
                report["warnings"].append(f"P[{i}]: keywords use comma — should use ；")

        # ── Punctuation ──
        total_alpha = sum(1 for c in text if c.isalpha())
        chinese_chars = sum(1 for c in text if is_chinese_char(c))
        is_cn = chinese_chars > total_alpha * 0.5 and chinese_chars > 5

        if is_cn:
            for en_punct, (cn_name, desc) in _EN_PUNCT_IN_CN_WARNING.items():
                if en_punct in text:
                    if en_punct == '.':
                        continue
                    for m in re.finditer(re.escape(en_punct), text):
                        pos = m.start()
                        pc = text[pos - 1] if pos > 0 else ''
                        nc = text[pos + 1] if pos + 1 < len(text) else ''
                        if is_chinese_char(pc) or is_chinese_char(nc):
                            punctuation_issues.append(
                                f"P[{i}]: {desc}: "
                                f"\"...{text[max(0,pos-5):pos+6]}...\" → {cn_name}"
                            )
                            break

            if '。' in text and expected_full_stop == '．':
                full_stop_issues.append(f"P[{i}]: '。'(U+3002) → expected {full_stop_name}")
            if '．' in text and expected_full_stop == '。':
                full_stop_issues.append(f"P[{i}]: '．'(U+FF0E) → expected {full_stop_name}")

            for m in re.finditer(r'(?<=[^\d])\.(?=\s|$)', text):
                pos = m.start()
                pc = text[pos - 1] if pos > 0 else ''
                if is_chinese_char(pc):
                    punctuation_issues.append(
                        f"P[{i}]: '.' after Chinese char → {full_stop_name}"
                    )
                    break

    # ── Aggregate ──
    if punctuation_issues:
        report["warnings"].extend(punctuation_issues)
    else:
        report["passed"].append("Punctuation OK (no English punct in Chinese text)")

    if full_stop_issues:
        report["manual_fixes"].extend(full_stop_issues)
    else:
        report["passed"].append(f"Full stop: {full_stop_name}")

    # ── Bilingual ──
    if bilingual:
        if not has_english_abstract:
            report["manual_fixes"].append("Missing English abstract (bilingual required)")
        if not has_english_keywords:
            report["manual_fixes"].append("Missing English keywords (bilingual required)")
        if has_chinese_refs and not has_english_refs:
            report["manual_fixes"].append("Chinese refs need English translations (bilingual)")
    else:
        report["passed"].append("Bilingual not required")

    # ── Abstract / Keywords format ──
    if has_abstract_inline:
        report["passed"].append("Abstract: inline format (【摘要】...)")
    else:
        report["manual_fixes"].append("Abstract should use: 【摘要】content")

    if has_keywords_inline:
        report["passed"].append("Keywords: inline format (【关键词】...)")
    else:
        report["manual_fixes"].append("Keywords should use: 【关键词】item1；item2；...")

    # ── Tables ──
    for ti, table in enumerate(doc.tables):
        report["manual_fixes"].append(
            f"Table {ti}: verify three-line style. "
            f"{'Check bilingual caption.' if bilingual else 'Check caption format.'}"
        )

    # ── References ──
    has_ref_section = any(
        '参考文献' in p.text or 'References' in p.text
        for p in doc.paragraphs
    )
    if has_ref_section:
        report["passed"].append("Reference section found")
    else:
        report["manual_fixes"].append("No 参考文献 section found")

    report["passed"].insert(0, f"Template: {template_name}")
    return report


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Verify formatted academic paper")
    parser.add_argument("document", help="Formatted .docx file")
    parser.add_argument("rules", help="Rules JSON file")
    parser.add_argument("--json", action="store_true", help="Output as JSON")
    args = parser.parse_args()

    with open(args.rules, 'r', encoding='utf-8') as f:
        rules = json.load(f)

    report = check_document(args.document, rules)

    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        print("=" * 60)
        print("FORMAT VERIFICATION REPORT")
        print("=" * 60)
        for section, symbol in [("passed", "[OK]"), ("warnings", "[WARN]"), ("manual_fixes", "[TODO]")]:
            items = report[section]
            label = {"passed": "PASSED", "warnings": "WARNINGS", "manual_fixes": "MANUAL FIXES NEEDED"}[section]
            print(f"\n[{label}] ({len(items)})")
            for item in items:
                print(f"  {symbol} {item}")
            if not items:
                print("  None")
        print("\n" + "=" * 60)


if __name__ == '__main__':
    main()
